"""导出 Markdown / Word（PRD 7.7）。pdf 后置。"""

from __future__ import annotations

import io
from typing import Any, Dict, List

from .models import MeetingSummary, Segment
from .templates_prompts import template_name
from .textproc import speaker_label


def to_markdown(meeting: Dict[str, Any], summary: MeetingSummary, segments: List[Segment]) -> str:
    dur = meeting.get("duration_sec") or 0
    mm, ss = divmod(int(dur), 60)
    hh, mm = divmod(mm, 60)
    L: List[str] = []
    L.append(f"# {summary.title}\n")
    L.append(
        f"> 会议类型：{template_name(meeting.get('template_type', 'general'))} ｜ "
        f"时长：{hh:02d}:{mm:02d}:{ss:02d} ｜ 原文件：{meeting.get('original_filename', '')}\n"
    )
    L.append("## 会议摘要\n")
    L.append(summary.summary + "\n")

    if summary.topics:
        L.append("## 关键讨论点\n")
        for t in summary.topics:
            L.append(f"- **{t.title}**（{t.source_time}）：{t.summary}")
        L.append("")

    if summary.decisions:
        L.append("## 已确认决策\n")
        for d in summary.decisions:
            L.append(f"- {d.content}（{d.source_time}）")
        L.append("")

    if summary.todos:
        L.append("## 待办事项\n")
        L.append("| 负责人 | 事项 | 截止时间 | 出处 |")
        L.append("| --- | --- | --- | --- |")
        for t in summary.todos:
            L.append(f"| {t.owner} | {t.task} | {t.deadline} | {t.source_time} |")
        L.append("")

    if summary.risks:
        L.append("## 风险问题\n")
        for r in summary.risks:
            L.append(f"- {r.content}（{r.source_time}）")
        L.append("")

    if summary.open_questions:
        L.append("## 未决问题\n")
        for q in summary.open_questions:
            L.append(f"- {q.content}（{q.source_time}）")
        L.append("")

    if segments:
        L.append("## 转写全文（证据）\n")
        for s in segments:
            L.append(f"- `[{s.start}]` **{speaker_label(s)}**：{s.text}")
        L.append("")

    return "\n".join(L)


def to_docx(meeting: Dict[str, Any], summary: MeetingSummary, segments: List[Segment]) -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading(summary.title, level=0)
    doc.add_paragraph(
        f"会议类型：{template_name(meeting.get('template_type', 'general'))} ｜ "
        f"原文件：{meeting.get('original_filename', '')}"
    )

    doc.add_heading("会议摘要", level=1)
    doc.add_paragraph(summary.summary)

    if summary.topics:
        doc.add_heading("关键讨论点", level=1)
        for t in summary.topics:
            doc.add_paragraph(f"{t.title}（{t.source_time}）：{t.summary}", style="List Bullet")

    if summary.decisions:
        doc.add_heading("已确认决策", level=1)
        for d in summary.decisions:
            doc.add_paragraph(f"{d.content}（{d.source_time}）", style="List Bullet")

    if summary.todos:
        doc.add_heading("待办事项", level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "负责人", "事项", "截止时间", "出处"
        for t in summary.todos:
            row = table.add_row().cells
            row[0].text, row[1].text, row[2].text, row[3].text = (
                t.owner, t.task, t.deadline, t.source_time
            )

    if summary.risks:
        doc.add_heading("风险问题", level=1)
        for r in summary.risks:
            doc.add_paragraph(f"{r.content}（{r.source_time}）", style="List Bullet")

    if summary.open_questions:
        doc.add_heading("未决问题", level=1)
        for q in summary.open_questions:
            doc.add_paragraph(f"{q.content}（{q.source_time}）", style="List Bullet")

    if segments:
        doc.add_heading("转写全文（证据）", level=1)
        for s in segments:
            doc.add_paragraph(f"[{s.start}] {speaker_label(s)}：{s.text}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
